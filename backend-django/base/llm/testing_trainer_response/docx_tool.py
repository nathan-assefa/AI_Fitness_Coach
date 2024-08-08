from docx import Document
from docx.shared import Pt
from typing import Dict
from io import BytesIO

def generate_docx(nutrition_data: Dict, chat_history: Dict) -> BytesIO:
    print(f"\n\n\n\n\n\n\n\n\n\n {nutrition_data} \n\n\n\n\n\n\n\n\n")
    print(f"\n\n\n\n\n\n\n\n\n\n {nutrition_data['days']} \n\nn\n\n\n\n\n\n\n")
    """
    This method generates a Docx file for the trainer to edit and forward as a PDF to the user.
    Args:
        nutrition_data (Dict): A dictionary containing the nutrition or workout plan data.
        chat_history (Dict): A dictionary containing the chat history data.

    Returns:
        BytesIO: The generated Docx file as a BytesIO object.
    """
    doc = Document()
    
    if not isinstance(nutrition_data, dict):
        raise TypeError("nutrition_data must be a dictionary")
    
    if nutrition_data["plan_type"] == "meal" or nutrition_data["plan_type"] == "meal_plan" or nutrition_data["plan_type"] == "meals":
        doc.add_heading(f'Nutrition Plan for {nutrition_data["month"]} {nutrition_data["year"]}', 0)
        
        for day_plan in nutrition_data['days']:
            doc.add_heading(f'Day {day_plan["day"]} - {day_plan["date"]}', level=1)
            
            for meal in day_plan['meals']:
                doc.add_heading(meal["type"], level=2)
                
                doc.add_heading('Items:', level=3)
                for item in meal["items"]:
                    doc.add_paragraph(f'- {item}', style='List Bullet')
                
                doc.add_heading('Nutritional Information:', level=3)
                nutritional_info = meal["nutritional_info"][0]
                doc.add_paragraph(f'Calories: {nutritional_info["calories"]} kcal')
                doc.add_paragraph(f'Protein: {nutritional_info["protein_grams"]} g')
                doc.add_paragraph(f'Carbohydrates: {nutritional_info["carbs_grams"]} g')
                doc.add_paragraph(f'Fat: {nutritional_info["fat_grams"]} g')
                
                doc.add_heading('Time:', level=3)
                doc.add_paragraph(f'{meal["time"]}')
                
            if day_plan.get("notes"):
                doc.add_heading('Notes:', level=2)
                doc.add_paragraph(day_plan["notes"])
                            
    else:
        doc.add_heading(f'Workout Plan for {nutrition_data["month"]} {nutrition_data["year"]}', 0)
        
        for day_plan in nutrition_data['days']:
            doc.add_heading(f'Day {day_plan["day"]} - {day_plan["date"]}', level=1)
            
            for workout in day_plan['workouts']:
                doc.add_heading(workout["type"].capitalize(), level=2)
                
                doc.add_heading('Description:', level=3)
                doc.add_paragraph(workout["description"])
                
                doc.add_heading('Duration:', level=3)
                doc.add_paragraph(f'{workout["duration_minutes"]} minutes')
                
                doc.add_heading('Intensity:', level=3)
                doc.add_paragraph(workout["intensity"].capitalize())
                
                if workout["equipment_needed"]:
                    doc.add_heading('Equipment Needed:', level=3)
                    for equipment in workout["equipment_needed"]:
                        doc.add_paragraph(f'- {equipment}', style='List Bullet')
                
                if workout["target_muscle_groups"]:
                    doc.add_heading('Target Muscle Groups:', level=3)
                    for muscle in workout["target_muscle_groups"]:
                        doc.add_paragraph(f'- {muscle}', style='List Bullet')
                
            if day_plan.get("notes"):
                doc.add_heading('Notes:', level=2)
                doc.add_paragraph(day_plan["notes"])
                  

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(20)

    for style in doc.styles:
        if style.type == 1: 
            style.font.name = 'Times New Roman'
            style.font.size = Pt(15)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream